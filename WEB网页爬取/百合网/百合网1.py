import requests
from lxml import etree
from docx import Document#docx文档
from docx.shared import Inches#插入图片
import time

session = requests.session()
document = Document()  # 创建docx文档
headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.25 Safari/537.36 Core/1.70.3756.400 QQBrowser/10.5.4039.400','Cookie': 'accessID=20200612155724877383; channel=qq-dh-k; code=140170-b; c=Pb62Ddki-1594168383887-4acef8a098b62-50251362; _fmdata=QVN3Ew47GDng3v84nBPVSEXPBAtzIj0DpeHfFy9nV8JdSTL186LcZf5AJHB6eq%2F3JDnrqM7GN9fZqY6TXUMj1HRPQXrZ6I8nqIQPTz74YA0%3D; _xid=evb0x%2BtvRq%2FOFgh3J5M8hzv02kQRSiUZeAnt59rb2hhnDF19V9yEy3OETftadvjYdT5FMeyVjLNX9lBZSdnwpg%3D%3D; cookie_pcc=1%7C%7Cqq-dh-k%7C%7C140170-b%7C%7Chttp%3A//daohang.qq.com/%3Ffr%3Dhmpage; Qs_lvt_336725=1591948645%2C1591948825%2C1594168365%2C1594213667; Qs_pv_336725=3949968046040140300%2C2591345869472359000%2C1580624541311368400%2C2707212606616901600; Hm_lvt_5497f3e5cf06014777f01fb0307e58f4=1591948645,1591948825,1594168384,1594213667; Hm_lpvt_5497f3e5cf06014777f01fb0307e58f4=1594213667; lastLoginDate=Wed%20Jul%2008%202020%2021%3A07%3A50%20GMT+0800%20%28%u4E2D%u56FD%u6807%u51C6%u65F6%u95F4%29; Hm_lvt_5caa30e0c191a1c525d4a6487bf45a9d=1594213672; AuthCookie=4BFFD62B611D896EF082FE1F7707ED5B06C2ACDE5DF1438B320A326458B02DA0C6E5940F3A828C4B904AF051C2A81FF4D7A6E2425F8EA858B437F9EDFEE40BE6D488D5ADA68FBC116BFB0F813A653583; AuthMsgCookie=E0A9E2736F2F4EA9926A80D126967B2EB037FCB858094067816F0597DD87851C801555B71BD28A4D1028A56E3D3E3EC4E8967372A2B0AAC788AF501D7E036C42CFED174701AFFA783E1B42F04DDA46D3; GCUserID=211695885; OnceLoginWEB=211695885; LoginEmail=17856889351%40mobile.baihe.com; userID=211695885; spmUserID=211695885; orderSource=10130301; tempID=5054555964; hasphoto=1; AuthCheckStatusCookie=C76E152D625A054C9E4B370B02BAD57B14D1D3015978705726A127B203C4931405E9D878DCC812E7; Hm_lpvt_5caa30e0c191a1c525d4a6487bf45a9d=1594214071; _fmdata=QVN3Ew47GDng3v84nBPVSEXPBAtzIj0DpeHfFy9nV8JdSTL186LcZf5AJHB6eq%2F3rYwskqAYyG11QbqqY%2BGpIxFybHTieBZ06My5aynsUy4%3D; noticeEvent_211695885=8; accessToken=baihe-1594214303587-8d41ffed46278'}
for i in range(1000000000):#哟个循环随机获取用户id
    url='https://profile1.baihe.com/?oppID='+str(i)#构造用户页面的url地址
    print(url)
    r=session.get(url,headers=headers)
    if r.status_code==200:#测试请求是否成功响应
        html=etree.HTML(r.content.decode('utf-8'))
        error=html.xpath('string(//div[@class="proBlank"]/dl/dd/p/a)')#获取错误的情况
        t='查看更多会员，请点击此处>'
        if error==t:
            continue
        else:#构造的用户页面成功后，爬取基本信息与图片
            girl=html.xpath('string(//div[@class="other"]/a/em)')[2:3]#确定女性
            if bool(girl)==True:
                try:
                    a2 = html.xpath('string(//div[@class="inter"]/p)')[:2]
                    #b2 = html.xpath('string(//div[@class="inter"]/p)')[4:7]
                    if int(a2) <= 27:
                        title = html.xpath('string(//div[@class="inter"]/p)')
                        a = str(title).replace('/', ',')
                        imgurl = html.xpath('string(//div[@class="big_pic"]/ul/li/a/img/@src)')  # 图片地址
                        if bool(imgurl) == True:  # 用布尔值确定图片地址，是否成功
                            url2 = 'http:' + imgurl
                            res = requests.get(url2, headers=headers)
                            with open('百合网/1/' + a + '.jpg', 'wb')as f:
                                f.write(res.content)

                            document.add_heading('userid: %d' % i, level=2)  # id标题
                            document.add_heading(title, level=2)  # 基本信息标题
                            document.add_picture('百合网/1/' + a + '.jpg', width=Inches(3.5), height=Inches(3.5))
                            paragraph = document.add_paragraph('_' * 30)
                            document.save('百合网.docx')
                            print('userid存在，id为： ' + str(i))
                            print(url)
                    else:
                        continue
                except:
                    continue
            else:
                continue

    else:
        continue



